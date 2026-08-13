"""Evaluation Script for PII Redaction Tool"""
import json
import sys
import os
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from src.detector import PIIDetector
from src.config import PII_CATEGORIES

class PIIEvaluator:
    """Evaluates PII detection performance"""
    
    def __init__(self, ground_truth_path: str):
        """Initialize evaluator with ground truth
        
        Args:
            ground_truth_path: Path to ground truth JSON file
        """
        with open(ground_truth_path, 'r') as f:
            self.ground_truth = json.load(f)
        
        self.detector = PIIDetector(use_spacy=False)
        self.fake_indicators = set()  # Track generated fake values to exclude from FP
    
    def evaluate_document(self, document_path: str) -> dict:
        """Evaluate PII detection on the ORIGINAL source document
        
        Args:
            document_path: Path to ORIGINAL DOCX document (not redacted)
            
        Returns:
            Dictionary containing evaluation metrics
        """
        # Load ORIGINAL document (input, not output)
        doc = Document(document_path)
        all_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += '\n' + cell.text
        
        # Detect PII in original document
        detections = self.detector.detect_all(all_text)
        
        # Calculate metrics per PII type
        results = {}
        
        for pii_type in PII_CATEGORIES:
            results[pii_type] = self._evaluate_pii_type(
                pii_type, 
                detections.get(pii_type, []),
                all_text
            )
        
        # Calculate overall metrics
        overall = self._calculate_overall_metrics(results)
        results['OVERALL'] = overall
        
        return results
    
    def _evaluate_pii_type(self, pii_type: str, detections: list, text: str) -> dict:
        """Evaluate detection for a specific PII type
        
        Args:
            pii_type: Type of PII
            detections: List of detected instances
            text: Full document text
            
        Returns:
            Dictionary with TP, FP, FN, precision, recall, F1
        """
        # Get ground truth for this type
        gt_items = [
            ann for ann in self.ground_truth.get('annotations', [])
            if ann['type'] == pii_type
        ]
        
        if not gt_items:
            # Check if this category is marked as not present
            not_present = self.ground_truth.get('categories_not_present', {})
            if pii_type in not_present:
                return {
                    'tp': 0,
                    'fp': 0,
                    'fn': 0,
                    'precision': 1.0,
                    'recall': 1.0,
                    'f1': 1.0,
                    'note': f'N/A - {not_present[pii_type]}',
                    'detected_count': len(set(d[0] for d in detections)),
                    'ground_truth_count': 0
                }
            else:
                # No ground truth annotations for this type
                return {
                    'tp': 0,
                    'fp': len(set(d[0] for d in detections)),
                    'fn': 0,
                    'precision': 0.0 if detections else 1.0,
                    'recall': 1.0,
                    'f1': 0.0,
                    'note': 'No ground truth annotations for this PII type',
                    'detected_count': len(set(d[0] for d in detections)),
                    'ground_truth_count': 0
                }
        
        # Extract detected values (normalize for comparison)
        detected_values = set(self._normalize_text(d[0]) for d in detections)
        ground_truth_values = set(self._normalize_text(gt['text']) for gt in gt_items)
        
        # Calculate TP, FP, FN using normalized matching
        tp_values = detected_values & ground_truth_values
        fp_values = detected_values - ground_truth_values
        fn_values = ground_truth_values - detected_values
        
        tp = len(tp_values)
        fp = len(fp_values)
        fn = len(fn_values)
        
        # Calculate metrics
        precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
        recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
        f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0.0
        
        return {
            'tp': tp,
            'fp': fp,
            'fn': fn,
            'precision': precision,
            'recall': recall,
            'f1': f1,
            'detected_count': len(detected_values),
            'ground_truth_count': len(ground_truth_values),
            'tp_examples': list(tp_values)[:5],
            'fp_examples': list(fp_values)[:5],
            'fn_examples': list(fn_values)[:5]
        }
    
    def _normalize_text(self, text: str) -> str:
        """Normalize text for comparison
        
        Args:
            text: Input text
            
        Returns:
            Normalized text
        """
        import re
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text.strip())
        # Lowercase for case-insensitive comparison
        text = text.lower()
        return text
    
    def _calculate_overall_metrics(self, results: dict) -> dict:
        """Calculate overall metrics across all PII types
        
        Args:
            results: Per-type results
            
        Returns:
            Overall metrics dictionary
        """
        total_tp = sum(r['tp'] for r in results.values() if 'tp' in r and 'note' not in r)
        total_fp = sum(r['fp'] for r in results.values() if 'fp' in r and 'note' not in r)
        total_fn = sum(r['fn'] for r in results.values() if 'fn' in r and 'note' not in r)
        
        precision = total_tp / (total_tp + total_fp) if (total_tp + total_fp) > 0 else 0.0
        recall = total_tp / (total_tp + total_fn) if (total_tp + total_fn) > 0 else 0.0
        f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0.0
        
        # Accuracy for PII detection: proportion of correct identifications
        # Accuracy = TP / (TP + FP + FN)
        accuracy = total_tp / (total_tp + total_fp + total_fn) if (total_tp + total_fp + total_fn) > 0 else 0.0
        
        return {
            'tp': total_tp,
            'fp': total_fp,
            'fn': total_fn,
            'precision': precision,
            'recall': recall,
            'f1': f1,
            'accuracy': accuracy
        }
    
    def generate_report(self, results: dict, output_path: str):
        """Generate evaluation report
        
        Args:
            results: Evaluation results
            output_path: Path to save report
        """
        report = []
        report.append("=" * 80)
        report.append("PII REDACTION EVALUATION REPORT")
        report.append("=" * 80)
        report.append("")
        report.append(f"Ground Truth Annotations: {self.ground_truth.get('total_annotations', 'N/A')}")
        report.append(f"Document: {self.ground_truth.get('document', 'Red Herring Prospectus')}")
        report.append("")
        
        # Per-type results
        report.append("PER-PII-TYPE RESULTS")
        report.append("=" * 80)
        report.append(f"{'PII Type':<15} {'GT':<6} {'Det':<6} {'TP':<6} {'FP':<6} {'FN':<6} {'Precision':<12} {'Recall':<12} {'F1':<12}")
        report.append("-" * 80)
        
        for pii_type in PII_CATEGORIES:
            if pii_type in results:
                r = results[pii_type]
                if 'note' in r:
                    report.append(f"{pii_type:<15} {r['note']}")
                else:
                    gt_count = r.get('ground_truth_count', 0)
                    det_count = r.get('detected_count', 0)
                    report.append(
                        f"{pii_type:<15} "
                        f"{gt_count:<6} "
                        f"{det_count:<6} "
                        f"{r['tp']:<6} "
                        f"{r['fp']:<6} "
                        f"{r['fn']:<6} "
                        f"{r['precision']:<12.2%} "
                        f"{r['recall']:<12.2%} "
                        f"{r['f1']:<12.2%}"
                    )
        
        report.append("-" * 80)
        
        # Overall results
        if 'OVERALL' in results:
            overall = results['OVERALL']
            report.append("")
            report.append("OVERALL METRICS")
            report.append("=" * 80)
            report.append(f"Total True Positives:  {overall['tp']}")
            report.append(f"Total False Positives: {overall['fp']}")
            report.append(f"Total False Negatives: {overall['fn']}")
            report.append("")
            report.append(f"Overall Precision: {overall['precision']:.2%}")
            report.append(f"Overall Recall:    {overall['recall']:.2%}")
            report.append(f"Overall F1 Score:  {overall['f1']:.2%}")
            report.append(f"Overall Accuracy:  {overall['accuracy']:.2%}")
            report.append("")
            report.append("Note: Accuracy = TP / (TP + FP + FN)")
        
        report.append("")
        report.append("=" * 80)
        
        # Add examples for categories with issues
        report.append("")
        report.append("EXAMPLE DETECTIONS")
        report.append("=" * 80)
        for pii_type in PII_CATEGORIES:
            if pii_type in results and 'note' not in results[pii_type]:
                r = results[pii_type]
                if r['tp'] > 0:
                    report.append(f"\n{pii_type} - True Positives (examples):")
                    for ex in r.get('tp_examples', [])[:3]:
                        report.append(f"  [TP] {ex}")
                if r['fp'] > 0 and r.get('fp_examples'):
                    report.append(f"\n{pii_type} - False Positives (examples):")
                    for ex in r.get('fp_examples', [])[:3]:
                        report.append(f"  [FP] {ex}")
                if r['fn'] > 0 and r.get('fn_examples'):
                    report.append(f"\n{pii_type} - False Negatives (examples):")
                    for ex in r.get('fn_examples', [])[:3]:
                        report.append(f"  [FN] {ex}")
        
        report.append("")
        report.append("=" * 80)
        
        # Write report
        report_text = '\n'.join(report)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_text)
        
        return report_text

def main():
    """Main evaluation function"""
    print("=" * 80)
    print("PII REDACTION EVALUATION")
    print("=" * 80)
    
    # Paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    ground_truth_path = os.path.join(script_dir, 'ground_truth.json')
    # IMPORTANT: Evaluate on ORIGINAL document, not redacted
    input_path = os.path.join(project_dir, 'input', 'Red Herring Prospectus.docx')
    output_report_path = os.path.join(script_dir, 'evaluation_results.txt')
    
    # Check files exist
    if not os.path.exists(ground_truth_path):
        print(f"Error: Ground truth file not found: {ground_truth_path}")
        sys.exit(1)
    
    if not os.path.exists(input_path):
        print(f"Error: Input document not found: {input_path}")
        sys.exit(1)
    
    # Run evaluation
    evaluator = PIIEvaluator(ground_truth_path)
    
    print(f"\nEvaluating ORIGINAL document: {input_path}")
    print(f"Ground truth annotations: {evaluator.ground_truth.get('total_annotations', 'N/A')}")
    print("This may take a moment...\n")
    
    results = evaluator.evaluate_document(input_path)
    
    # Generate and display report
    report = evaluator.generate_report(results, output_report_path)
    print(report)
    
    print(f"\nReport saved to: {output_report_path}")
    print("=" * 80)

if __name__ == '__main__':
    main()
