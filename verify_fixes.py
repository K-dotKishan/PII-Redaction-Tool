"""Verify all three critical fixes were applied"""
from docx import Document

print("="*70)
print("VERIFICATION OF CRITICAL FIXES")
print("="*70)

orig = Document('input/Red Herring Prospectus.docx')
redact = Document('output/redacted_prospectus.docx')

orig_text = ' '.join([p.text for p in orig.paragraphs])
redact_text = ' '.join([p.text for p in redact.paragraphs])

# ISSUE #1: Verify promoter/executive names were redacted (RECALL)
print("\n[ISSUE #1] Promoter & Executive Names - Redaction Check:")
print("-"*70)
promoter_names = [
    'Kushal Subbayya Hegde',
    'Pushpa Kushal Hegde',
    'Rajesh Kushal Hegde',
    'Rohit Kushal Hegde',
    'Sarthak Malvadkar',
    'Sandesh Bhagwat',
    'Amod Joshi',
    'Dinesh Hirachand Munot',
    'Ajay Shriram Patil'
]

total_orig = 0
total_redact = 0
for name in promoter_names:
    orig_count = orig_text.count(name)
    redact_count = redact_text.count(name)
    total_orig += orig_count
    total_redact += redact_count
    if orig_count > 0:
        status = "✓ REDACTED" if redact_count == 0 else "✗ LEAKED"
        print(f"{name:30} Orig: {orig_count:2}  Redact: {redact_count:2}  {status}")

print(f"\n{'TOTAL':30} Orig: {total_orig:2}  Redact: {total_redact:2}")
recall_pass = total_redact == 0 and total_orig > 0
print(f"{'ISSUE #1 STATUS: ' + ('✓ PASSED' if recall_pass else '✗ FAILED'):>70}")

# ISSUE #2: Verify legal/generic terms NOT over-redacted (PRECISION)
print("\n[ISSUE #2] Legal & Generic Terms - Preservation Check:")
print("-"*70)
protected_terms = [
    'Companies Act',
    'SEBI',
    'Risk Factors',
    'Our Company',
    'The Offer',
    'Restated Financial',
    'Net Worth',
    'ICICI Securities',
    'HDFC Bank',
]

precision_ok = True
for term in protected_terms:
    orig_count = orig_text.count(term)
    redact_count = redact_text.count(term)
    if orig_count > 0:
        preservation_rate = (redact_count / orig_count * 100) if orig_count > 0 else 0
        status = "✓ PRESERVED" if preservation_rate >= 90 else "✗ OVER-REDACTED"
        if preservation_rate < 90:
            precision_ok = False
        print(f"{term:25} Orig: {orig_count:3}  Redact: {redact_count:3}  {preservation_rate:5.1f}%  {status}")

print(f"{'ISSUE #2 STATUS: ' + ('✓ PASSED' if precision_ok else '✗ FAILED'):>70}")

# ISSUE #3: Verify embedded images were redacted
print("\n[ISSUE #3] Embedded Images - Redaction Check:")
print("-"*70)

# Count image redaction placeholders
image_placeholders = redact_text.count('[IMAGE_CONTAINING_PII_REDACTED]')
print(f"Image redaction placeholders found: {image_placeholders}")

# Try to count inline shapes in original vs redacted
try:
    orig_images = 0
    for para in orig.paragraphs:
        if hasattr(para, '_element'):
            orig_images += len(para._element.xpath('.//w:drawing'))
    
    redact_images = 0
    for para in redact.paragraphs:
        if hasattr(para, '_element'):
            redact_images += len(para._element.xpath('.//w:drawing'))
    
    images_removed = orig_images - redact_images
    print(f"Original document images: {orig_images}")
    print(f"Redacted document images: {redact_images}")
    print(f"Images removed: {images_removed}")
    
    image_pass = images_removed > 0 or image_placeholders > 0
    print(f"{'ISSUE #3 STATUS: ' + ('✓ PASSED' if image_pass else '✗ FAILED'):>70}")
except Exception as e:
    print(f"Could not verify image removal: {e}")
    print(f"{'ISSUE #3 STATUS: ' + ('✓ PASSED (placeholders found)' if image_placeholders > 0 else '? UNKNOWN'):>70}")

# Overall summary
print("\n" + "="*70)
print("OVERALL SUMMARY")
print("="*70)
print(f"ISSUE #1 (Recall - Name Detection):     {'✓ PASSED' if recall_pass else '✗ FAILED'}")
print(f"ISSUE #2 (Precision - Over-redaction):  {'✓ PASSED' if precision_ok else '✗ FAILED'}")
print(f"ISSUE #3 (Image Redaction):             {'✓ PASSED' if image_placeholders > 0 else '? CHECK'}")
print("="*70)
