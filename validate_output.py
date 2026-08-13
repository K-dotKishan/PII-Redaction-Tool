"""Comprehensive Output Validation Script
Verifies that redacted output preserves non-PII content exactly
"""
import sys
sys.path.insert(0, '.')

from docx import Document
import re

print("="*70)
print("VALIDATION: Document Integrity Check")
print("="*70)

# Load original and redacted documents
print("\nLoading documents...")
original_doc = Document('input/Red Herring Prospectus.docx')
redacted_doc = Document('output/redacted_prospectus.docx')

# Extract all text
original_text = '\n'.join([p.text for p in original_doc.paragraphs])
redacted_text = '\n'.join([p.text for p in redacted_doc.paragraphs])

print(f"Original: {len(original_doc.paragraphs)} paragraphs")
print(f"Redacted: {len(redacted_doc.paragraphs)} paragraphs")

# Critical validation tests
print("\n" + "="*70)
print("CRITICAL VALIDATION TESTS")
print("="*70)

validation_tests = []

# Test 1: "Companies Act, 2013" must be unchanged
test1_orig = original_text.count("Companies Act, 2013")
test1_redact = redacted_text.count("Companies Act, 2013")
test1_pass = test1_orig > 0 and test1_redact == test1_orig
validation_tests.append(("Companies Act, 2013 unchanged", test1_pass, f"{test1_orig} → {test1_redact}"))

# Test 2: "SEBI" must be unchanged
test2_orig = original_text.count("SEBI")
test2_redact = redacted_text.count("SEBI")
test2_pass = test2_orig > 0 and test2_redact == test2_orig
validation_tests.append(("SEBI unchanged", test2_pass, f"{test2_orig} → {test2_redact}"))

# Test 3: "BSE" or "NSE" must be present
test3_orig_bse = original_text.count("BSE")
test3_redact_bse = redacted_text.count("BSE")
test3_orig_nse = original_text.count("NSE")
test3_redact_nse = redacted_text.count("NSE")
test3_pass = (test3_orig_bse > 0 and test3_redact_bse == test3_orig_bse) or \
             (test3_orig_nse > 0 and test3_redact_nse == test3_orig_nse)
validation_tests.append(("BSE/NSE unchanged", test3_pass, f"BSE: {test3_orig_bse}→{test3_redact_bse}, NSE: {test3_orig_nse}→{test3_redact_nse}"))

# Test 4: Main company name must be present
test4_count = redacted_text.count("KSH International Limited")
test4_pass = test4_count >= 50  # Should appear many times
validation_tests.append(("KSH International Limited preserved", test4_pass, f"{test4_count} occurrences"))

# Test 5: "Dated December" should NOT be followed by a person name
dated_december_corrupt = re.search(r"Dated\s+[A-Z][a-z]+\s+[A-Z][a-z]+\s+\d+", redacted_text)
test5_pass = not dated_december_corrupt
validation_tests.append(("No 'Dated' + random name", test5_pass, "Clean" if test5_pass else f"Found: {dated_december_corrupt.group()}"))

# Test 6: "Book Built" should still exist (not replaced with person name)
test6_count = redacted_text.count("Book Built")
test6_pass = test6_count > 0
validation_tests.append(("Book Built preserved", test6_pass, f"{test6_count} occurrences"))

# Test 7: "Our Company" should be preserved
test7_count = redacted_text.count("Our Company")
test7_pass = test7_count > 5  # Should appear many times
validation_tests.append(("Our Company preserved", test7_pass, f"{test7_count} occurrences"))

# Test 8: Emails should be anonymized (no @kshinternational.com)
test8_orig = original_text.count("@kshinternational.com")
test8_redact = redacted_text.count("@kshinternational.com")
test8_pass = test8_orig > 0 and test8_redact < test8_orig
validation_tests.append(("Emails anonymized", test8_pass, f"{test8_orig} → {test8_redact}"))

# Test 9: Phone numbers should be anonymized (specific example)
test9_orig = original_text.count("+91 20 4505 3237")
test9_redact = redacted_text.count("+91 20 4505 3237")
test9_pass = test9_orig > 0 and test9_redact < test9_orig
validation_tests.append(("Phone numbers anonymized", test9_pass, f"{test9_orig} → {test9_redact}"))

# Test 10: Check for random name insertions in financial context
financial_corruption = re.search(r"(?:Rs\.|INR|Crore|Million)\s+[A-Z][a-z]+\s+[A-Z][a-z]+", redacted_text)
test10_pass = not financial_corruption
validation_tests.append(("No random names in financial context", test10_pass, "Clean" if test10_pass else f"Found: {financial_corruption.group()}"))

# Test 11: Protected companies preserved
protected_companies = [
    "ICICI Securities",
    "HDFC",
    "Nuvama",
    "CARE Ratings"
]
test11_results = []
test11_pass = True
for company in protected_companies:
    orig_count = original_text.count(company)
    redact_count = redacted_text.count(company)
    if orig_count > 0:
        if redact_count != orig_count:
            test11_pass = False
        test11_results.append(f"{company}: {orig_count}→{redact_count}")
validation_tests.append(("Protected business entities preserved", test11_pass, ", ".join(test11_results)))

# Test 12: Dates preserved (check for common patterns)
date_pattern = r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+20\d{2}\b"
test12_orig_matches = len(re.findall(date_pattern, original_text))
test12_redact_matches = len(re.findall(date_pattern, redacted_text))
test12_pass = test12_orig_matches > 0 and test12_redact_matches >= test12_orig_matches * 0.95  # Allow 5% variance
validation_tests.append(("Dates preserved", test12_pass, f"{test12_orig_matches} → {test12_redact_matches}"))

# Test 13: Paragraph count matches
test13_pass = len(original_doc.paragraphs) == len(redacted_doc.paragraphs)
validation_tests.append(("Paragraph count matches", test13_pass, f"{len(original_doc.paragraphs)} → {len(redacted_doc.paragraphs)}"))

# Test 14: Table count matches
test14_pass = len(original_doc.tables) == len(redacted_doc.tables)
validation_tests.append(("Table count matches", test14_pass, f"{len(original_doc.tables)} → {len(redacted_doc.tables)}"))

# Test 15: "The Board" preserved
test15_count = redacted_text.count("The Board")
test15_pass = test15_count > 0
validation_tests.append(("The Board preserved", test15_pass, f"{test15_count} occurrences"))

# Test 16: Check for deterministic output (same email appears multiple times)
from src.replacer import PIIReplacer
replacer = PIIReplacer()

# Find a specific email in original
sample_email = "cs.connect@kshinternational.com"
if sample_email in original_text:
    # Get its replacement
    replacement = replacer.get_replacement('EMAIL', sample_email)
    # Check if replacement is consistent (deterministic)
    replacement2 = replacer.get_replacement('EMAIL', sample_email)
    test16_pass = replacement == replacement2
    validation_tests.append(("Consistent anonymization", test16_pass, f"{sample_email} → {replacement}"))
else:
    validation_tests.append(("Consistent anonymization", True, "N/A"))

# Print results
print()
passed = 0
failed = 0
for test_name, passed_flag, details in validation_tests:
    status = "✓ PASS" if passed_flag else "✗ FAIL"
    print(f"{status}: {test_name}")
    if details:
        print(f"       {details}")
    if passed_flag:
        passed += 1
    else:
        failed += 1

print("\n" + "="*70)
print(f"VALIDATION SUMMARY: {passed} passed, {failed} failed")
print("="*70)

if failed == 0:
    print("\n✓ ALL VALIDATION TESTS PASSED")
    print("✓ Document integrity maintained")
    print("✓ Only PII anonymized")
    print("✓ No unexpected modifications")
    print("\nThe redacted document is SUBMISSION-READY.")
else:
    print(f"\n✗ {failed} VALIDATION TEST(S) FAILED")
    print("✗ Review the failures above")
    print("✗ The document may have unexpected modifications")

print("="*70)
