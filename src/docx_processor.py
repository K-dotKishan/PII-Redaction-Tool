"""DOCX Processing Module"""
from docx import Document
from docx.shared import Inches
from typing import Dict, List, Tuple
from .detector import PIIDetector
from .replacer import PIIReplacer
import io

class DocxProcessor:
    """Processes DOCX files for PII redaction"""
    
    def __init__(self, use_spacy: bool = True):
        """Initialize the DOCX processor
        
        Args:
            use_spacy: Whether to use spaCy for NER
        """
        self.detector = PIIDetector(use_spacy=use_spacy)
        self.replacer = PIIReplacer()
        self.image_redaction_count = 0
    
    def process_document(self, input_path: str, output_path: str) -> Dict:
        """Process a DOCX document and redact PII
        
        Args:
            input_path: Path to input DOCX file
            output_path: Path to save redacted DOCX file
            
        Returns:
            Dictionary containing processing statistics
        """
        print(f"Loading document: {input_path}")
        doc = Document(input_path)
        
        print("Detecting PII in paragraphs...")
        all_detections = {}
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                detections = self.detector.detect_all(para.text)
                self._merge_detections(all_detections, detections)
                
                # Replace PII in paragraph
                if any(detections.values()):
                    redacted_text, _ = self.replacer.replace_in_text(para.text, detections)
                    self._replace_paragraph_text(para, redacted_text)
        
        # ISSUE #3 FIX: Process and redact embedded images (PAN/Aadhaar cards, etc.)
        print("Scanning for embedded images (PII documents)...")
        self._redact_embedded_images(doc)
        
        # Process tables
        print("Detecting PII in tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            detections = self.detector.detect_all(para.text)
                            self._merge_detections(all_detections, detections)
                            
                            # Replace PII in cell
                            if any(detections.values()):
                                redacted_text, _ = self.replacer.replace_in_text(para.text, detections)
                                self._replace_paragraph_text(para, redacted_text)
        
        # Save redacted document
        print(f"Saving redacted document: {output_path}")
        doc.save(output_path)
        
        # Get statistics
        stats = self.replacer.get_statistics(all_detections)
        
        # Add image redaction count
        if self.image_redaction_count > 0:
            stats['IMAGES_REDACTED'] = self.image_redaction_count
        
        return {
            'detections': all_detections,
            'statistics': stats,
            'replacement_map': self.replacer.replacement_map
        }
    
    def _redact_embedded_images(self, doc):
        """Redact embedded images that may contain PII (ISSUE #3 FIX)
        
        Replaces inline shapes (embedded images) with a text placeholder
        indicating that an image containing PII has been redacted.
        
        Args:
            doc: python-docx Document object
        """
        redacted_count = 0
        
        # Iterate through all paragraphs to find inline shapes
        for para in doc.paragraphs:
            # Check if paragraph contains inline shapes (embedded images)
            if hasattr(para, '_element') and hasattr(para._element, 'xpath'):
                # Find all drawing elements in the paragraph
                drawings = para._element.xpath('.//w:drawing')
                
                if drawings:
                    # This paragraph contains embedded images
                    # Replace images with placeholder text
                    for drawing in drawings:
                        try:
                            # Remove the drawing element
                            drawing.getparent().remove(drawing)
                            redacted_count += 1
                        except:
                            pass
                    
                    # Add placeholder text if paragraph is now empty
                    if not para.text.strip() or para.text == '':
                        para.text = '[IMAGE_CONTAINING_PII_REDACTED]'
                        # Make it visible
                        if para.runs:
                            para.runs[0].bold = True
        
        # Also check tables for embedded images
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if hasattr(para, '_element') and hasattr(para._element, 'xpath'):
                            drawings = para._element.xpath('.//w:drawing')
                            
                            if drawings:
                                for drawing in drawings:
                                    try:
                                        drawing.getparent().remove(drawing)
                                        redacted_count += 1
                                    except:
                                        pass
                                
                                if not para.text.strip() or para.text == '':
                                    para.text = '[IMAGE_CONTAINING_PII_REDACTED]'
                                    if para.runs:
                                        para.runs[0].bold = True
        
        self.image_redaction_count = redacted_count
        
        if redacted_count > 0:
            print(f"✓ Redacted {redacted_count} embedded image(s) (potential PII documents)")
    
    def _replace_paragraph_text(self, paragraph, new_text: str):
        """Replace paragraph text while preserving formatting
        
        Args:
            paragraph: python-docx paragraph object
            new_text: New text to set
        """
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''
        
        # Add new text in first run (or create new run)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
    
    def _merge_detections(self, all_detections: Dict, new_detections: Dict):
        """Merge new detections into all_detections
        
        Args:
            all_detections: Accumulated detections
            new_detections: New detections to merge
        """
        for pii_type, matches in new_detections.items():
            if pii_type not in all_detections:
                all_detections[pii_type] = []
            all_detections[pii_type].extend(matches)
    
    def verify_redaction(self, output_path: str) -> Dict:
        """Verify that redaction was successful
        
        Args:
            output_path: Path to redacted DOCX file
            
        Returns:
            Dictionary containing verification results
        """
        print(f"\nVerifying redaction: {output_path}")
        
        try:
            doc = Document(output_path)
            
            # Extract all text
            all_text = '\n'.join([para.text for para in doc.paragraphs])
            
            # Scan for remaining PII
            remaining_pii = self.detector.detect_all(all_text)
            
            # Count remaining PII
            remaining_counts = {}
            for pii_type, matches in remaining_pii.items():
                unique_matches = set(match[0] for match in matches)
                if unique_matches:
                    remaining_counts[pii_type] = len(unique_matches)
            
            verification_passed = len(remaining_counts) == 0
            
            return {
                'success': True,
                'file_readable': True,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'remaining_pii': remaining_counts,
                'verification_passed': verification_passed
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file_readable': False
            }
