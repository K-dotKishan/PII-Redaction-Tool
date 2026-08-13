"""Check where HDFC Bank Limited appears and if it's protected"""
from docx import Document

# Check original
orig_doc = Document('input/Red Herring Prospectus.docx')
redact_doc = Document('output/redacted_prospectus.docx')

print("="*70)
print("HDFC Bank Limited - Investigation")
print("="*70)

# Check in paragraphs
print("\n1. In Paragraphs:")
for i, p in enumerate(orig_doc.paragraphs):
    if 'HDFC Bank Limited' in p.text:
        print(f"  Para {i}: {p.text[:100]}")
        print(f"  Redacted: {redact_doc.paragraphs[i].text[:100]}")
        print()

# Check in tables
print("\n2. In Tables:")
for i, table in enumerate(orig_doc.tables):
    for j, row in enumerate(table.rows):
        for k, cell in enumerate(row.cells):
            if 'HDFC Bank Limited' in cell.text:
                print(f"  Table {i}, Row {j}, Cell {k}:")
                print(f"    Original: {cell.text[:150]}")
                redact_cell = redact_doc.tables[i].rows[j].cells[k]
                print(f"    Redacted: {redact_cell.text[:150]}")
                print()

# Test the protection logic
print("\n3. Testing protection logic:")
import sys
sys.path.insert(0, '.')
from src.detector import PIIDetector

detector = PIIDetector(use_spacy=False)

test_text = "HDFC Bank Limited"
print(f"  Test: '{test_text}'")
print(f"  Is protected: {detector._is_protected_company(test_text)}")

test_text2 = "HDFC Bank"  
print(f"  Test: '{test_text2}'")
print(f"  Is protected: {detector._is_protected_company(test_text2)}")

# Test detection
test_text3 = "Bankers: HDFC Bank Limited"
companies = detector.detect_companies_pattern(test_text3)
print(f"\n  Detection test: '{test_text3}'")
print(f"  Companies detected: {companies}")
